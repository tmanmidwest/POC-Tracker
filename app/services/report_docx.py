"""Server-side Word (.docx) rendering for the use-case library and project report.

Mirrors the PDF reports' content as an *editable* document, built with
python-docx (pure Python — no system libraries, so this works everywhere the
WeasyPrint PDF path can't). Category/use-case names are real Word headings so
the native Table of Contents field (added near the top) can build itself.
"""

from __future__ import annotations

import io
import logging
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from app.models import Project
from app.services import screenshots as screenshot_store

log = logging.getLogger(__name__)

_MAX_IMG_WIDTH = Inches(6.0)


def _accent(color: str | None) -> RGBColor:
    raw = (color or "#1e293b").lstrip("#")
    if len(raw) != 6:
        raw = "1e293b"
    return RGBColor(int(raw[0:2], 16), int(raw[2:4], 16), int(raw[4:6], 16))


def _enable_update_fields(doc: Document) -> None:
    """Tell Word to refresh fields (the TOC) when the document is opened."""
    settings = doc.settings.element
    upd = OxmlElement("w:updateFields")
    upd.set(qn("w:val"), "true")
    settings.append(upd)


def _add_toc(doc: Document, levels: str = "1-3") -> None:
    """Insert a native Word Table of Contents field for the given heading levels."""
    p = doc.add_paragraph()
    run = p.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f'TOC \\o "{levels}" \\h \\z \\u'
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click and choose “Update Field” to build the contents."
    sep.append(placeholder)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    r = run._r
    for el in (begin, instr, sep, end):
        r.append(el)


def _title_block(doc: Document, brand: str, title: str, subtitle: str, color: str | None) -> None:
    brand_p = doc.add_paragraph()
    brand_run = brand_p.add_run(brand)
    brand_run.bold = True
    brand_run.font.size = Pt(10)
    brand_run.font.color.rgb = _accent(color)
    brand_p.paragraph_format.space_after = Pt(0)

    h = doc.add_heading(title, level=0)
    for r in h.runs:
        r.font.color.rgb = _accent(color)

    sub = doc.add_paragraph(subtitle)
    sub.runs[0].font.size = Pt(9)
    sub.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)


def _labeled(doc: Document, label: str, text: str) -> None:
    """A small uppercase label followed by the value on its own line."""
    p = doc.add_paragraph()
    lab = p.add_run(label.upper())
    lab.bold = True
    lab.font.size = Pt(8)
    lab.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    p.paragraph_format.space_after = Pt(0)
    body = doc.add_paragraph(text)
    body.paragraph_format.space_after = Pt(6)


def _to_bytes(doc: Document) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Library report
# ---------------------------------------------------------------------------


def build_library_docx(context: dict[str, Any]) -> bytes:
    """Editable Word version of the use-case library share document."""
    library = context["library"]
    groups = context["groups"]
    brand = context["branding"]
    doc = Document()

    sub = (
        f"Use Case Library · {context['total']} use case"
        f"{'' if context['total'] == 1 else 's'}"
    )
    if library.description:
        sub += f" · {library.description}"
    sub += f" · Generated {context['generated_on']}"
    _title_block(doc, brand.get("name", ""), library.name, sub, brand.get("color"))

    doc.add_heading("Contents", level=1)
    _add_toc(doc, levels="1-2")

    for g in groups:
        doc.add_heading(g["category"], level=1)
        for e in g["entries"]:
            name = e.name
            if e.default_reference_number:
                name = f"{e.default_reference_number} · {name}"
            doc.add_heading(name, level=2)
            if e.feature_type:
                doc.add_paragraph(f"Feature: {e.feature_type.name}")
            if e.description:
                _labeled(doc, "Description", e.description)
            if e.success_validation:
                _labeled(doc, "Success validation", e.success_validation)

    if not groups:
        doc.add_paragraph("This library has no active use cases.")

    _enable_update_fields(doc)
    return _to_bytes(doc)


# ---------------------------------------------------------------------------
# Project report
# ---------------------------------------------------------------------------


def _kv_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light Grid Accent 1"
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value or "—"
        for p in cells[0].paragraphs:
            for r in p.runs:
                r.bold = True


def _embed_screenshots(doc: Document, use_case) -> None:
    if not use_case.screenshots:
        return
    _labeled_label = doc.add_paragraph()
    lab = _labeled_label.add_run("SCREENSHOTS")
    lab.bold = True
    lab.font.size = Pt(8)
    lab.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    for shot in use_case.screenshots:
        try:
            path = screenshot_store.path_for(shot)
            doc.add_picture(str(path), width=_MAX_IMG_WIDTH)
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.LEFT
        except Exception:  # missing file or format python-docx can't embed
            log.warning("docx_screenshot_skipped", extra={"screenshot_id": shot.id})
            doc.add_paragraph("[screenshot unavailable]")
            continue
        if shot.caption:
            cap = doc.add_paragraph(shot.caption)
            cap.runs[0].font.size = Pt(8)
            cap.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)


def build_project_report_docx(project: Project, context: dict[str, Any]) -> bytes:
    """Editable Word version of the full project report (with screenshots)."""
    brand = context["branding"]
    progress = context["progress"]
    doc = Document()

    subtitle = f"{project.customer.name} · {project.status.name}"
    if project.is_archived:
        subtitle += " · Archived"
    _title_block(
        doc, brand.get("name", ""), project.display_name, subtitle, brand.get("color")
    )

    doc.add_heading("Contents", level=1)
    _add_toc(doc, levels="1-3")

    if project.exec_summary:
        doc.add_heading("Executive summary", level=1)
        doc.add_paragraph(project.exec_summary)

    doc.add_heading("Project details", level=1)
    _kv_table(
        doc,
        [
            ("Status", project.status.name),
            ("Customer", project.customer.name),
            (
                "Use-case progress",
                f"{progress['done']}/{progress['total']} complete ({progress['pct']}%)",
            ),
            ("POC name", project.name or "—"),
            (
                "Sales Engineer",
                project.sales_engineer.display_label if project.sales_engineer else "—",
            ),
            ("Account Executive", project.account_executive or "—"),
            (
                "Start date",
                project.start_date.strftime("%b %d, %Y") if project.start_date else "—",
            ),
            (
                "End date",
                project.end_date.strftime("%b %d, %Y") if project.end_date else "—",
            ),
            ("Salesforce opportunity", project.salesforce_opp_url or "—"),
        ],
    )

    if project.customer.contacts:
        doc.add_heading("Contacts", level=1)
        for c in project.customer.contacts:
            role = c.role.name if c.role else "—"
            doc.add_paragraph(f"{c.name} ({role}) · {c.email or '—'} · {c.phone or '—'}")

    doc.add_heading("Use cases", level=1)
    groups = context["use_case_groups"]
    if not groups:
        doc.add_paragraph("No use cases.")
    for group in groups:
        doc.add_heading(group["category"], level=2)
        for uc in group["use_cases"]:
            name = uc.name
            if uc.reference_number:
                name = f"{uc.reference_number} · {name}"
            doc.add_heading(name, level=3)
            status = uc.status.name if uc.status else "—"
            feature = uc.feature_type.name if uc.feature_type else "—"
            completed = uc.completed_on.strftime("%b %d, %Y") if uc.completed_on else "—"
            meta = doc.add_paragraph(
                f"Status: {status}  ·  Feature: {feature}  ·  "
                f"Source: {uc.source}  ·  Completed: {completed}"
            )
            meta.runs[0].font.size = Pt(9)
            meta.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
            if uc.description:
                _labeled(doc, "Description", uc.description)
            if uc.success_validation:
                _labeled(doc, "Success validation", uc.success_validation)
            if uc.comments:
                _labeled(doc, "Comments", uc.comments)
            _embed_screenshots(doc, uc)

    _enable_update_fields(doc)
    return _to_bytes(doc)
